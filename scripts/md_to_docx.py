#!/usr/bin/env python3
"""Convert BLOG.md into a clean .docx for publishing.

Handles: H1/H2/H3, paragraphs, bullet lists, ordered lists, blockquotes,
fenced code blocks, inline **bold**, *italic*, `code`, and image placeholders
(rendered as italic captions). Mermaid blocks are rendered as a labeled
placeholder paragraph because Word cannot render mermaid.
"""

import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


INLINE_PATTERN = re.compile(
    r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|\[[^\]]+\]\([^\)]+\))"
)


def add_inline_runs(paragraph, text: str):
    """Parse inline markdown (bold/italic/code/links) into runs."""
    parts = INLINE_PATTERN.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(10)
        elif part.startswith("[") and "](" in part:
            # Link — render the link text only, styled.
            label = part[1 : part.index("]")]
            run = paragraph.add_run(label)
            run.font.color.rgb = RGBColor(0x1F, 0x6F, 0xEB)
            run.underline = True
        else:
            paragraph.add_run(part)


def convert(md_path: Path, out_path: Path) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()

    doc = Document()

    # Global defaults
    normal = doc.styles["Normal"]
    normal.font.name = "Georgia"
    normal.font.size = Pt(11)

    i = 0
    in_code = False
    code_lang = ""
    code_buf: list[str] = []

    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()

        # Fenced code block
        if line.startswith("```"):
            if not in_code:
                in_code = True
                code_lang = line[3:].strip()
                code_buf = []
            else:
                # Close block
                in_code = False
                if code_lang.lower() == "mermaid":
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run(
                        "[mermaid flow chart — renders on web; see figure above for the prose equivalent]"
                    )
                    run.italic = True
                    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                else:
                    p = doc.add_paragraph()
                    for cl in code_buf:
                        run = p.add_run(cl + "\n")
                        run.font.name = "Courier New"
                        run.font.size = Pt(9)
                code_buf = []
            i += 1
            continue

        if in_code:
            code_buf.append(raw)
            i += 1
            continue

        # Horizontal rule
        if line.strip() == "---":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("• • •")
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            i += 1
            continue

        # Empty line
        if not line.strip():
            i += 1
            continue

        # Headings
        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            heading = doc.add_heading(level=min(level, 4))
            heading.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
            )
            run = heading.add_run(text)
            run.font.name = "Georgia"
            if level == 1:
                run.font.size = Pt(22)
            elif level == 2:
                run.font.size = Pt(16)
            else:
                run.font.size = Pt(13)
            i += 1
            continue

        # Image lines -> italic caption
        m = re.match(r"^\s*!\[([^\]]*)\]\(([^\)]+)\)\s*$", line)
        if m:
            alt = m.group(1)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"[figure: {alt}]")
            run.italic = True
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1
            continue

        # Blockquote
        if line.startswith(">"):
            buf = []
            while i < len(lines) and lines[i].startswith(">"):
                buf.append(lines[i].lstrip(">").strip())
                i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            run = p.add_run(" ".join(buf))
            run.italic = True
            run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            continue

        # Bullet list
        if re.match(r"^\s*[-*]\s+", line):
            while i < len(lines) and re.match(r"^\s*[-*]\s+", lines[i]):
                text = re.sub(r"^\s*[-*]\s+", "", lines[i])
                p = doc.add_paragraph(style="List Bullet")
                add_inline_runs(p, text)
                i += 1
            continue

        # Ordered list
        if re.match(r"^\s*\d+\.\s+", line):
            while i < len(lines) and re.match(r"^\s*\d+\.\s+", lines[i]):
                text = re.sub(r"^\s*\d+\.\s+", "", lines[i])
                p = doc.add_paragraph(style="List Number")
                add_inline_runs(p, text)
                i += 1
            continue

        # Default: paragraph (with continuation lines)
        buf = [line]
        i += 1
        while (
            i < len(lines)
            and lines[i].strip()
            and not lines[i].startswith(("#", ">", "!", "```"))
            and not re.match(r"^\s*[-*]\s+", lines[i])
            and not re.match(r"^\s*\d+\.\s+", lines[i])
            and lines[i].strip() != "---"
        ):
            buf.append(lines[i].rstrip())
            i += 1
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(10)
        add_inline_runs(p, " ".join(buf))

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    print(f"wrote {out_path}")


if __name__ == "__main__":
    src = Path(sys.argv[1]) if len(sys.argv) > 1 else Path("BLOG.md")
    dst = (
        Path(sys.argv[2])
        if len(sys.argv) > 2
        else Path.home() / "Downloads" / "on-the-architecture-of-remembering.docx"
    )
    convert(src, dst)
